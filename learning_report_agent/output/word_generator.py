"""Word (.docx) report generator — produces branded Word documents from report dicts."""
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_BLUE  = RGBColor(0x00, 0x78, 0xD4)
_DARK  = RGBColor(0x1a, 0x1a, 0x2e)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREY  = RGBColor(0xF5, 0xF5, 0xF5)


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = _BLUE
    run.font.size = Pt(14 if level == 1 else 11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _metric_row(doc: Document, metrics: list[tuple]):
    """Add a single-row table as metric band: [(value, label), ...]"""
    tbl = doc.add_table(rows=2, cols=len(metrics))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, (val, lbl) in enumerate(metrics):
        # value cell
        vc = tbl.rows[0].cells[i]
        _set_cell_bg(vc, "EBF5FB")
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(str(val))
        vr.bold = True
        vr.font.size = Pt(18)
        vr.font.color.rgb = _BLUE
        # label cell
        lc = tbl.rows[1].cells[i]
        _set_cell_bg(lc, "EBF5FB")
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(lbl)
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()


def _data_table(doc: Document, headers: list, rows: list):
    if not rows:
        doc.add_paragraph("No data available.")
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "0078D4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _WHITE
        run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else "F5F5F5"
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run("—" if (val is None or str(val) == "") else str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()


def generate_word(report: dict) -> bytes:
    """Convert a report dict to a branded .docx. Returns raw bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(report.get("title", "Learning Report"))
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _BLUE

    meta_p = doc.add_paragraph(
        f"Generated: {str(report.get('generated_at', ''))[:19]}  |  "
        f"Type: {report.get('report_type', '').replace('_', ' ').title()}  |  "
        f"LTM – HackToFuture 2026"
    )
    meta_p.runs[0].font.size = Pt(9)
    meta_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("─" * 80)

    rtype = report.get("report_type")
    if   rtype == "compliance":    _compliance(doc, report)
    elif rtype == "learning_kpis": _kpis(doc, report)
    elif rtype == "skill_gap":     _skill_gap(doc, report)
    else:
        import json
        doc.add_paragraph(json.dumps(report, default=str, indent=2)[:3000])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _compliance(doc: Document, report: dict):
    _heading(doc, "Executive Summary")
    _metric_row(doc, [
        (f"{report.get('avg_compliance_rate', 0)}%", "Avg Compliance Rate"),
        (len(report.get("at_risk_departments", [])),  "At-Risk Departments"),
        (report.get("total_at_risk", 0),              "Overdue Learners"),
    ])

    summary = report.get("summary", [])
    if summary:
        _heading(doc, "Completion Rate by Department & Course", level=2)
        _data_table(doc,
            ["Department", "Course", "Completion %", "Avg Score", "Enrolled", "Completed"],
            [[r.get("department"), r.get("course_title"),
              f"{r.get('completion_rate', 0)}%", r.get("avg_score") or "N/A",
              r.get("total_enrolled"), r.get("completed")]
             for r in summary[:35]]
        )

    overdue = report.get("overdue_learners", [])
    if overdue:
        _heading(doc, f"Overdue Learners ({len(overdue)} records)", level=2)
        _data_table(doc,
            ["Employee", "Department", "Role", "Course", "Status"],
            [[r.get("name"), r.get("department", ""), r.get("role", ""),
              r.get("course_title"), r.get("status")]
             for r in overdue[:30]]
        )


def _kpis(doc: Document, report: dict):
    top = report.get("top_performing", {})
    _heading(doc, "Executive Summary")
    _metric_row(doc, [
        (f"{report.get('org_avg_completion', 0)}%", "Org Avg Completion"),
        (f"{report.get('org_avg_hours', 0)}h",       "Avg Hours / Employee"),
        (top.get("department", "N/A"),                "Top Department"),
    ])

    kpis = report.get("kpis", [])
    if kpis:
        _heading(doc, "Department Learning KPIs", level=2)
        _data_table(doc,
            ["Department", "HC", "Completion %", "Required %",
             "Avg Hrs", "Budget ($)", "Utilised %"],
            [[r.get("department"), r.get("headcount"),
              f"{r.get('completion_rate', 0)}%", f"{r.get('required_completion_rate', 0)}%",
              f"{r.get('avg_hours_per_employee', 0)}h",
              f"${r.get('budget_usd', 0):,.0f}", f"{r.get('budget_utilization', 0)}%"]
             for r in kpis]
        )


def _skill_gap(doc: Document, report: dict):
    _heading(doc, "Overview")
    _metric_row(doc, [
        (report.get("total_skill_areas", 0),       "Skill Areas Analysed"),
        (len(report.get("overrated_skills", [])),  "Overrated Skills"),
        (len(report.get("underrated_skills", [])), "Underrated Skills"),
    ])

    gaps = report.get("gaps", [])
    if gaps:
        _heading(doc, "Skill Gap Detail (Self − Manager Rating)", level=2)
        _data_table(doc,
            ["Department", "Skill", "Self", "Manager", "Gap", "Employees"],
            [[r.get("department"), r.get("skill"),
              f"{r.get('avg_self_rating', 0):.2f}",
              f"{r.get('avg_manager_rating'):.2f}" if r.get("avg_manager_rating") else "N/A",
              f"{r.get('gap', 0):.2f}", r.get("employee_count")]
             for r in gaps[:35]]
        )

    for section_key, section_title in [
        ("overrated_skills",  "Top Overrated Skills (self >> manager)"),
        ("underrated_skills", "Top Underrated Skills (manager >> self)"),
    ]:
        items = report.get(section_key, [])
        if items:
            _heading(doc, section_title, level=2)
            _data_table(doc,
                ["Department", "Skill", "Gap"],
                [[r.get("department"), r.get("skill"), f"{r.get('gap', 0):.2f}"]
                 for r in items]
            )
